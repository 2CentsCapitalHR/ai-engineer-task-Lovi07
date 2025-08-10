import os
import re
import json
import streamlit as st
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from openai import OpenAI
from red_flag_detector import check_placeholders, check_missing_clauses, llm_detect_wrong_jurisdiction
from process_checklists import PROCESS_CHECKLISTS

# ====== CONFIG ======
os.environ["OPENAI_API_KEY"] = "Open AI api Key"
TEMPLATE_FOLDER = "templates"
openai_client = OpenAI()

# ====== FUNCTIONS ======
def extract_text_from_docx(path):
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def build_template_index(template_folder=TEMPLATE_FOLDER):
    docs, metadata = [], []
    for file_name in os.listdir(template_folder):
        if file_name.lower().endswith(".docx"):
            path = os.path.join(template_folder, file_name)
            text = extract_text_from_docx(path)
            docs.append(text)
            metadata.append({"title": file_name})
    embeddings = OpenAIEmbeddings()
    return FAISS.from_texts(docs, embeddings, metadatas=metadata)

# ====== Annotation helper ======
def highlight_and_comment_docx(input_path, issues, output_path):
    """Highlight matching text in yellow and add inline COMMENT runs."""
    doc = Document(input_path)
    paras = list(doc.paragraphs)

    def _highlight_in_paragraph(paragraph, match_text, issue_text):
        for run in paragraph.runs:
            if match_text in run.text:
                before, after = run.text.split(match_text, 1)
                run.text = before
                match_run = paragraph.add_run(match_text)
                try:
                    match_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                except Exception:
                    match_run.bold = True
                comment_run = paragraph.add_run(f"  <<COMMENT: {issue_text}>>")
                comment_run.italic = True
                if after:
                    paragraph.add_run(after)
                return True
        return False

    not_found_issues = []
    for issue in issues:
        snippet = issue.get("snippet") or issue.get("section", "")
        found = False
        if snippet:
            for p in paras:
                if snippet.lower() in p.text.lower():
                    m = re.search(re.escape(snippet.strip()), p.text, re.IGNORECASE)
                    if m:
                        actual_text = p.text[m.start():m.end()]
                        found = _highlight_in_paragraph(p, actual_text, issue.get("issue", "Issue found"))
                        if found:
                            break
        if not found:
            not_found_issues.append(issue)

    if not_found_issues:
        doc.add_page_break()
        try:
            doc.add_heading("Review Comments (Auto-generated)", level=1)
        except KeyError:
            # If Heading 1 is missing, use normal paragraph with bold text
            para = doc.add_paragraph()
            run = para.add_run("Review Comments (Auto-generated)")
            run.bold = True
            run.font.size = doc.styles["Normal"].font.size

        for issue in not_found_issues:
            para = doc.add_paragraph()
            para.add_run("Document: ").bold = True
            para.add_run(str(issue.get("document", "")))
            para.add_run("\nIssue: ").bold = True
            para.add_run(str(issue.get("issue", "")))
            para.add_run("\nSuggestion: ").bold = True
            para.add_run(str(issue.get("suggestion", "")))


    doc.save(output_path)
    return output_path

# ====== BUILD INDEX ======
if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = build_template_index()

# ====== STREAMLIT APP ======
st.set_page_config(page_title="ADGM Document Checker", layout="wide")
st.title("📄 Identify, Review & Process Detection")

uploaded_files = st.file_uploader("Upload .docx files", type=["docx"], accept_multiple_files=True)

summary_report = {
    "process": None,
    "documents_uploaded": 0,
    "required_documents": 0,
    "missing_document": [],
    "issues_found": []
}

if uploaded_files:
    uploaded_doc_types = []

    for uploaded_file in uploaded_files:
        temp_path = f"temp_{uploaded_file.name}"
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        uploaded_text = extract_text_from_docx(temp_path)

        # Identify document type
        results = st.session_state.vectorstore.similarity_search(uploaded_text, k=1)
        best_match = results[0]
        doc_type = best_match.metadata["title"]
        uploaded_doc_types.append(doc_type)

        # ---- Checks ----
        placeholders = check_placeholders(temp_path, doc_type)
        missing_clauses = check_missing_clauses(temp_path, doc_type)

        if placeholders:
            for p in placeholders:
                summary_report["issues_found"].append({
                    "document": uploaded_file.name,
                    "issue": f"Unfilled placeholder: {p}",
                    "severity": "High",
                    "suggestion": "Fill this placeholder with the correct value",
                    "snippet": p
                })

        if missing_clauses:
            for c in missing_clauses:
                summary_report["issues_found"].append({
                    "document": uploaded_file.name,
                    "issue": f"Missing clause: {c}",
                    "severity": "High",
                    "suggestion": "Add the missing clause to match ADGM requirements",
                    "snippet": c
                })

        # ====== LLM-based Wrong Jurisdiction Detection ======
        llm_issues = llm_detect_wrong_jurisdiction(
            doc_text=uploaded_text,
            doc_name=doc_type,
            model="gpt-5",
            openai_client=openai_client
        )

        for issue in llm_issues:
            summary_report["issues_found"].append({
                "document": uploaded_file.name,
                "issue": issue.get("issue", ""),
                "severity": issue.get("severity", "Medium"),
                "suggestion": issue.get("suggestion", "N/A"),
                "snippet": issue.get("snippet", "")  # now will highlight wrong jurisdiction phrase
            })

        # ====== Annotate and save reviewed DOCX ======
        file_issues = [it for it in summary_report["issues_found"] if it["document"] == uploaded_file.name]
        reviewed_path = f"reviewed_{uploaded_file.name}"
        highlight_and_comment_docx(temp_path, file_issues, reviewed_path)

        with open(reviewed_path, "rb") as f:
            st.download_button(
                label=f"⬇ Download reviewed file: {uploaded_file.name}",
                data=f,
                file_name=reviewed_path,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # ====== Detect Process ======
    for process, required_docs in PROCESS_CHECKLISTS.items():
        matched_count = sum(1 for d in required_docs if d in uploaded_doc_types)
        if matched_count > 0:
            summary_report["process"] = process
            summary_report["documents_uploaded"] = matched_count
            summary_report["required_documents"] = len(required_docs)
            summary_report["missing_document"] = [d for d in required_docs if d not in uploaded_doc_types]
            break

    # ====== Show Summary ======
    st.subheader("📊 Summary Report")
    st.json(summary_report)
    # ====== Download JSON Report ======
    json_bytes = json.dumps(summary_report, indent=4).encode('utf-8')
    st.download_button(
        label="📥 Download JSON Report",
        data=json_bytes,
        file_name="compliance_report.json",
        mime="application/json"
    )
